"""
Utilities for extracting textual content from user-provided documents.

The goal is not to provide perfect fidelity but to offer a unified way to
inline file contents when forwarding them to providers that do not natively
support binary uploads.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional


MAX_EXTRACTED_CHARS = 6000


@dataclass
class ExtractedDocument:
    """Represents extracted data from a user supplied file."""

    name: str
    mime_type: str
    text: str
    truncated: bool = False

    def as_prompt_block(self) -> str:
        """Format the extracted content for inclusion in a chat prompt."""

        header = f"[파일: {self.name} | 타입: {self.mime_type}]"
        body = self.text
        if self.truncated:
            body += "\n... (추가 내용 생략됨)"
        return f"{header}\n{body.strip()}"


def _truncate_text(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS], True


def _read_text_file(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    truncated_text, truncated = _truncate_text(text)
    return ExtractedDocument(
        name=path.name,
        mime_type="text/plain" if path.suffix.lower() != ".md" else "text/markdown",
        text=truncated_text,
        truncated=truncated,
    )


def _read_markdown_file(path: Path) -> ExtractedDocument:
    return _read_text_file(path)


def _read_pdf(path: Path) -> ExtractedDocument:
    try:
        from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        pages_text = []
        for page in reader.pages:
            try:
                pages_text.append(page.extract_text() or "")
            except Exception:
                continue
        raw_text = "\n".join(pages_text)
        if not raw_text.strip():
            raw_text = "(PDF에서 텍스트를 추출하지 못했습니다. 요약이 필요하면 원문을 참조하세요.)"
    except Exception:
        raw_text = "(PyPDF2가 설치되어 있지 않거나 PDF 파싱에 실패했습니다. 문서를 직접 확인해주세요.)"

    truncated_text, truncated = _truncate_text(raw_text)
    return ExtractedDocument(
        name=path.name,
        mime_type="application/pdf",
        text=truncated_text,
        truncated=truncated,
    )


def _read_docx(path: Path) -> ExtractedDocument:
    try:
        import docx  # type: ignore

        document = docx.Document(str(path))
        paragraphs = [para.text for para in document.paragraphs]
        raw_text = "\n".join(paragraphs)
        if not raw_text.strip():
            raw_text = "(DOCX 파일에서 텍스트를 찾을 수 없습니다. 문서를 확인해주세요.)"
    except Exception:
        raw_text = "(python-docx가 설치되어 있지 않거나 DOCX 파싱에 실패했습니다. 문서를 직접 확인해주세요.)"

    truncated_text, truncated = _truncate_text(raw_text)
    return ExtractedDocument(
        name=path.name,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        text=truncated_text,
        truncated=truncated,
    )


def extract_document(path: Path) -> Optional[ExtractedDocument]:
    """Extract readable text from a supported document.

    Returns ``None`` when the file type is not currently supported.
    """

    suffix = path.suffix.lower()
    if suffix in {".txt"}:
        return _read_text_file(path)
    if suffix in {".md"}:
        return _read_markdown_file(path)
    if suffix in {".pdf"}:
        return _read_pdf(path)
    if suffix in {".docx"}:
        return _read_docx(path)
    # Legacy .doc files are not well supported without extra dependencies.
    if suffix == ".doc":
        placeholder = "(구형 DOC 형식은 직접 텍스트로 변환해 주세요. 현재 모듈은 DOCX를 우선 지원합니다.)"
        truncated_text, truncated = _truncate_text(placeholder)
        return ExtractedDocument(
            name=path.name,
            mime_type="application/msword",
            text=truncated_text,
            truncated=truncated,
        )
    return None


